"""
Simplified Word document processor for PII detection and redaction.
Focuses on core functionality while maintaining document structure.
"""

import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Tuple, Any, Optional
from datetime import datetime

from docx import Document
from docx.shared import RGBColor

from ..core.pii_detector import PIIDetector
from .base_processor import BaseFileProcessor


class WordProcessor(BaseFileProcessor):
    """
    Simplified Word document processor that extracts text and redacts PII.
    Maintains basic document structure while focusing on functionality.
    """
    
    def __init__(self, file_path: str = None):
        # Handle both factory pattern (no file_path) and legacy pattern (with file_path)
        if file_path:
            super().__init__(file_path)
        
        self.logger = logging.getLogger(__name__)
        self.pii_detector = PIIDetector()
        self.logger.info("Word processor initialized")
    
    def extract_text(self) -> Dict[str, Any]:
        """
        Extract text from Word document for legacy API compatibility.
        
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            if not hasattr(self, 'file_path') or not self.file_path:
                raise ValueError("File path not set for Word processor")
            
            # Simple Word document text extraction using python-docx
            from docx import Document
            
            doc = Document(self.file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                text_content += "\n"
            
            # Detect PII in the extracted text
            pii_findings = self.pii_detector.detect_pii(text_content)
            
            # Calculate PII statistics
            entity_types = {}
            for finding in pii_findings:
                pii_type = finding.get('entity_type', 'unknown')
                entity_types[pii_type] = entity_types.get(pii_type, 0) + 1
            
            return {
                'text': text_content,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'file_type': 'word'
                },
                'confidence': 0.9,  # High confidence for Word text extraction
                'pii_findings': pii_findings,
                'total_pii_count': len(pii_findings),
                'entity_types': entity_types
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document: {str(e)}")
            return {
                'text': '',
                'metadata': {'error': str(e)},
                'confidence': 0.0
            }
    
    def process_file(self, file_path: str, output_dir: str) -> Dict[str, Any]:
        """
        Process Word document with PII detection and redaction.
        
        Args:
            file_path: Path to the input Word document
            output_dir: Directory to save the redacted document
            
        Returns:
            Dictionary containing processing results and redacted file info
        """
        try:
            self.logger.info(f"Processing Word document: {file_path}")
            
            # Extract text from Word document
            full_text, paragraph_texts = self._extract_text_from_docx(file_path)
            
            if not full_text.strip():
                self.logger.warning("No text found in Word document")
                return {
                    'total_entities': 0,
                    'entity_types': {},
                    'redacted_file_path': None,
                    'original_text_length': 0,
                    'redacted_text_length': 0,
                    'processing_status': 'completed',
                    'confidence_score': 1.0,
                    'pages_processed': 1
                }
            
            # Detect PII in the document
            pii_findings = self.pii_detector.detect_pii(full_text)
            
            if not pii_findings:
                self.logger.info("No PII detected in Word document")
                return {
                    'total_entities': 0,
                    'entity_types': {},
                    'redacted_file_path': None,
                    'original_text_length': len(full_text),
                    'redacted_text_length': len(full_text),
                    'processing_status': 'completed',
                    'confidence_score': 1.0,
                    'pages_processed': 1
                }
            
            # Calculate PII statistics
            entity_types = {}
            for finding in pii_findings:
                pii_type = finding.get('entity_type', 'unknown')
                entity_types[pii_type] = entity_types.get(pii_type, 0) + 1
            
            # Create redacted Word document
            redacted_file_path = self._create_redacted_docx(
                file_path, paragraph_texts, pii_findings, output_dir
            )
            
            # Calculate processing metrics
            processing_result = {
                'total_entities': len(pii_findings),
                'entity_types': entity_types,
                'redacted_file_path': redacted_file_path,
                'original_text_length': len(full_text),
                'redacted_text_length': len(full_text),  # For now, keep original length
                'processing_status': 'completed',
                'confidence_score': self._calculate_confidence_score(pii_findings),
                'pages_processed': 1,
                'paragraphs_processed': len(paragraph_texts)
            }
            
            self.logger.info(f"Word document processing completed. Found {len(pii_findings)} PII entities")
            return processing_result
            
        except Exception as e:
            self.logger.error(f"Error processing Word document {file_path}: {str(e)}")
            raise
    
    def _extract_text_from_docx(self, file_path: str) -> Tuple[str, List[str]]:
        """
        Extract text from Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Tuple of (full_text, list_of_paragraph_texts)
        """
        try:
            doc = Document(file_path)
            paragraph_texts = []
            
            for paragraph in doc.paragraphs:
                paragraph_texts.append(paragraph.text)
            
            # Also extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    table_texts.append(" | ".join(row_text))
            
            # Combine all text
            all_texts = paragraph_texts + table_texts
            full_text = "\n".join(all_texts)
            
            self.logger.debug(f"Extracted text from {len(paragraph_texts)} paragraphs and {len(table_texts)} table rows")
            return full_text, paragraph_texts
            
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document: {str(e)}")
            raise
    
    def _create_redacted_docx(
        self, 
        original_path: str, 
        paragraph_texts: List[str], 
        pii_findings: List[Dict[str, Any]], 
        output_dir: str
    ) -> str:
        """
        Create redacted Word document with PII entities replaced.
        
        Args:
            original_path: Path to original Word document
            paragraph_texts: List of paragraph texts
            pii_findings: List of detected PII entities
            output_dir: Directory to save redacted document
            
        Returns:
            Path to the redacted Word document
        """
        try:
            # Generate output filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            original_name = Path(original_path).stem
            output_filename = f"{original_name}_REDACTED_{timestamp}.docx"
            output_path = Path(output_dir) / output_filename
            
            # Open original document
            original_doc = Document(original_path)
            
            # Create new document for redacted content
            redacted_doc = Document()
            
            # Copy document properties
            self._copy_document_properties(original_doc, redacted_doc)
            
            # Process paragraphs with PII redaction
            self._process_paragraphs(
                original_doc, redacted_doc, pii_findings
            )
            
            # Process tables with PII redaction
            self._process_tables(
                original_doc, redacted_doc, pii_findings
            )
            
            # Save the redacted document
            redacted_doc.save(str(output_path))
            
            self.logger.info(f"Redacted Word document saved to: {output_path}")
            return str(output_path)
            
        except Exception as e:
            self.logger.error(f"Error creating redacted Word document: {str(e)}")
            raise
    
    def _copy_document_properties(self, source_doc: Document, target_doc: Document) -> None:
        """
        Copy document properties from source to target.
        
        Args:
            source_doc: Original Word document
            target_doc: New redacted document
        """
        try:
            # Copy core properties if available
            if hasattr(source_doc, 'core_properties'):
                if source_doc.core_properties.title:
                    target_doc.core_properties.title = source_doc.core_properties.title
                if source_doc.core_properties.author:
                    target_doc.core_properties.author = source_doc.core_properties.author
                if source_doc.core_properties.subject:
                    target_doc.core_properties.subject = source_doc.core_properties.subject
            
            self.logger.debug("Document properties copied successfully")
            
        except Exception as e:
            self.logger.warning(f"Could not copy all document properties: {str(e)}")
    
    def _process_paragraphs(
        self, 
        original_doc: Document, 
        redacted_doc: Document, 
        pii_entities: List[Dict[str, Any]]
    ) -> None:
        """
        Process document paragraphs and apply PII redaction.
        
        Args:
            original_doc: Original Word document
            redacted_doc: New redacted document
            pii_entities: List of detected PII entities
        """
        try:
            # Process each paragraph in the original document
            for paragraph in original_doc.paragraphs:
                if not paragraph.text.strip():
                    # Add empty paragraph to maintain structure
                    redacted_doc.add_paragraph()
                    continue
                
                # Check if paragraph contains PII
                contains_pii = self._text_contains_pii(paragraph.text, pii_entities)
                
                if contains_pii:
                    # Redact the paragraph
                    redacted_text = self._redact_text(paragraph.text, pii_entities)
                    new_para = redacted_doc.add_paragraph(redacted_text)
                    
                    # Copy paragraph formatting
                    self._copy_paragraph_formatting(paragraph, new_para)
                    
                else:
                    # Keep original text and formatting
                    new_para = redacted_doc.add_paragraph(paragraph.text)
                    self._copy_paragraph_formatting(paragraph, new_para)
                        
        except Exception as e:
            self.logger.error(f"Error processing paragraphs: {str(e)}")
            raise
    
    def _process_tables(
        self, 
        original_doc: Document, 
        redacted_doc: Document, 
        pii_entities: List[Dict[str, Any]]
    ) -> None:
        """
        Process tables in the document with PII redaction.
        
        Args:
            original_doc: Original Word document
            redacted_doc: New redacted document
            pii_entities: List of PII entities to redact
        """
        try:
            for table in original_doc.tables:
                # Create new table in redacted document
                new_table = redacted_doc.add_table(
                    rows=len(table.rows), 
                    cols=len(table.columns)
                )
                
                # Process each cell
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text
                        
                        # Check if cell contains PII
                        if self._text_contains_pii(cell_text, pii_entities):
                            # Redact cell content
                            redacted_cell_text = self._redact_text(cell_text, pii_entities)
                            new_table.cell(row_idx, col_idx).text = redacted_cell_text
                        else:
                            # Keep original content
                            new_table.cell(row_idx, col_idx).text = cell_text
                            
        except Exception as e:
            self.logger.warning(f"Error processing tables: {str(e)}")
    
    def _text_contains_pii(
        self, 
        text: str, 
        pii_entities: List[Dict[str, Any]]
    ) -> bool:
        """
        Check if text contains any PII entities.
        
        Args:
            text: Text content
            pii_entities: List of detected PII entities
            
        Returns:
            True if text contains PII, False otherwise
        """
        for entity in pii_entities:
            if entity['text'] in text:
                return True
        return False
    
    def _redact_text(
        self, 
        text: str, 
        pii_entities: List[Dict[str, Any]]
    ) -> str:
        """
        Redact PII entities in text.
        
        Args:
            text: Original text
            pii_entities: List of PII entities to redact
            
        Returns:
            Text with PII entities redacted
        """
        redacted_text = text
        
        # Find entities that appear in this text
        relevant_entities = [
            entity for entity in pii_entities 
            if entity['text'] in text
        ]
        
        for entity in relevant_entities:
            # Create professional redaction using block characters
            entity_type = entity['entity_type'].replace('_', ' ').title()
            # Use block characters that match the length of original text
            redacted_length = len(entity['text'])
            block_chars = '█' * min(redacted_length, 20)  # Limit to 20 chars max
            redaction = f"🔒[{entity_type}:{block_chars}]"
            redacted_text = redacted_text.replace(entity['text'], redaction)
        
        return redacted_text
    
    def _copy_paragraph_formatting(
        self, 
        source_para: Any, 
        target_para: Any
    ) -> None:
        """
        Copy formatting from source paragraph to target paragraph.
        
        Args:
            source_para: Original paragraph with formatting
            target_para: New paragraph to apply formatting to
        """
        try:
            # Copy paragraph alignment
            if hasattr(source_para, 'alignment') and source_para.alignment:
                target_para.alignment = source_para.alignment
            
            # Copy paragraph style if available
            if hasattr(source_para, 'style') and source_para.style:
                target_para.style = source_para.style
                
            # Copy run-level formatting for the first run
            if source_para.runs and target_para.runs:
                source_run = source_para.runs[0]
                target_run = target_para.runs[0]
                
                if hasattr(source_run, 'bold'):
                    target_run.bold = source_run.bold
                if hasattr(source_run, 'italic'):
                    target_run.italic = source_run.italic
                if hasattr(source_run, 'underline'):
                    target_run.underline = source_run.underline
                
                if hasattr(source_run, 'font'):
                    if source_run.font.size:
                        target_run.font.size = source_run.font.size
                    if source_run.font.name:
                        target_run.font.name = source_run.font.name
                    
        except Exception as e:
            self.logger.warning(f"Could not copy all paragraph formatting: {str(e)}")
    
    def _calculate_confidence_score(self, pii_findings: List[Dict[str, Any]]) -> float:
        """
        Calculate confidence score based on PII detection results.
        
        Args:
            pii_findings: List of PII detection findings
            
        Returns:
            Confidence score between 0.0 and 1.0
        """
        # Base confidence on successful processing
        base_confidence = 0.85
        
        # Boost confidence if entities were found and processed
        if len(pii_findings) > 0:
            base_confidence = 0.95
        
        return base_confidence
    
    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        return ['.docx', '.doc']
    
    def validate_file(self, file_path: str) -> bool:
        """
        Validate that the file is a readable Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            True if file is valid, False otherwise
        """
        try:
            doc = Document(file_path)
            # Basic validation - check if document has content or structure
            return len(doc.paragraphs) >= 0
        except Exception as e:
            self.logger.error(f"Invalid Word document {file_path}: {str(e)}")
            return False